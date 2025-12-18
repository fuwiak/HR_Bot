"""
Script for indexing knowledge base documents into Qdrant
Поддерживает Word, Excel, PDF файлы
Использует OpenAI API (text-embedding-3-small) для генерации эмбеддингов
"""
import os
import sys
import logging
import asyncio
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
import hashlib

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
log = logging.getLogger()

# Импорты
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams, PointStruct
    QDRANT_AVAILABLE = True
except ImportError:
    QDRANT_AVAILABLE = False
    log.error("❌ qdrant-client не установлен. Установите: pip install qdrant-client")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log.warning("⚠️ python-docx не установлен. Установите: pip install python-docx")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    log.warning("⚠️ openpyxl не установлен. Установите: pip install openpyxl")

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    log.warning("⚠️ pypdf не установлен. Установите: pip install pypdf")

from qdrant_helper import (
    get_qdrant_client,
    generate_embedding,
    ensure_collection,
    COLLECTION_NAME,
    EMBEDDING_DIMENSION
)

# ===================== DOCUMENT PARSING =====================

def parse_docx(file_path: Path) -> Dict:
    """Парсинг Word документа"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx не установлен"}
    
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        text = "\n".join(text_parts)
        return {
            "text": text,
            "title": doc.core_properties.title or file_path.stem,
            "type": "docx",
            "pages": len([p for p in doc.paragraphs if p.text.strip()])
        }
    except Exception as e:
        log.error(f"❌ Ошибка парсинга DOCX {file_path}: {e}")
        return {"error": str(e)}

def parse_excel(file_path: Path) -> Dict:
    """Парсинг Excel файла"""
    if not EXCEL_AVAILABLE:
        return {"error": "openpyxl не установлен"}
    
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"\n--- Лист: {sheet_name} ---\n")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)
        
        text = "\n".join(text_parts)
        return {
            "text": text,
            "title": file_path.stem,
            "type": "excel",
            "sheets": workbook.sheetnames
        }
    except Exception as e:
        log.error(f"❌ Ошибка парсинга Excel {file_path}: {e}")
        return {"error": str(e)}

def parse_pdf(file_path: Path) -> Dict:
    """Парсинг PDF файла"""
    if not PDF_AVAILABLE:
        return {"error": "pypdf не установлен"}
    
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text.strip())
        
        text = "\n".join(text_parts)
        return {
            "text": text,
            "title": file_path.stem,
            "type": "pdf",
            "pages": len(reader.pages)
        }
    except Exception as e:
        log.error(f"❌ Ошибка парсинга PDF {file_path}: {e}")
        return {"error": str(e)}

# ===================== CHUNKING =====================

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Разбиение текста на чанки для индексации
    
    Args:
        text: Текст для разбиения
        chunk_size: Размер чанка (символов)
        overlap: Перекрытие между чанками (символов)
    
    Returns:
        Список чанков
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Пытаемся разбить по предложениям для более осмысленных чанков
        if end < len(text):
            # Ищем последнее предложение в чанке
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            split_point = max(last_period, last_newline)
            
            if split_point > start + chunk_size * 0.5:  # Если нашли разумное место для разбиения
                chunk = chunk[:split_point + 1]
                end = start + split_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap  # Перекрытие
        
        if start >= len(text):
            break
    
    return chunks

# ===================== INDEXING =====================

def index_document(
    file_path: Path,
    category: str = "другое",
    metadata: Optional[Dict] = None
) -> bool:
    """
    Индексирование одного документа в Qdrant
    
    Args:
        file_path: Путь к файлу
        category: Категория документа (кейсы, методики, шаблоны, статьи)
        metadata: Дополнительные метаданные
    
    Returns:
        True при успехе, False при ошибке
    """
    client = get_qdrant_client()
    if not client:
        log.error("❌ Qdrant клиент недоступен")
        return False
    
    # Определяем тип файла и парсим
    suffix = file_path.suffix.lower()
    
    if suffix == ".docx":
        doc_data = parse_docx(file_path)
    elif suffix in [".xlsx", ".xls"]:
        doc_data = parse_excel(file_path)
    elif suffix == ".pdf":
        doc_data = parse_pdf(file_path)
    else:
        log.warning(f"⚠️ Неподдерживаемый формат файла: {suffix}")
        return False
    
    if "error" in doc_data:
        log.error(f"❌ Ошибка парсинга {file_path}: {doc_data['error']}")
        return False
    
    text = doc_data.get("text", "")
    if not text.strip():
        log.warning(f"⚠️ Пустой документ: {file_path}")
        return False
    
    # Разбиваем на чанки
    chunks = chunk_text(text, chunk_size=1000, overlap=200)
    log.info(f"📄 Документ {file_path.name}: {len(chunks)} чанков")
    
    # Создаем коллекцию если нужно
    if not ensure_collection():
        log.error("❌ Не удалось создать коллекцию")
        return False
    
    # Индексируем каждый чанк
    points = []
    for i, chunk in enumerate(chunks):
        # Генерируем эмбеддинг
        embedding = generate_embedding(chunk)
        if embedding is None:
            log.warning(f"⚠️ Не удалось сгенерировать эмбеддинг для чанка {i+1}")
            continue
        
        # Создаем уникальный ID
        doc_id = hashlib.md5(f"{file_path}_{i}".encode()).hexdigest()
        point_id = int(doc_id[:8], 16)
        
        # Payload
        payload = {
            "document_title": doc_data.get("title", file_path.stem),
            "file_path": str(file_path),
            "file_name": file_path.name,
            "category": category,
            "chunk_index": i,
            "total_chunks": len(chunks),
            "text": chunk[:500],  # Первые 500 символов для предпросмотра
            "content": chunk,  # Полный текст
            "type": doc_data.get("type", "unknown"),
            "indexed_at": datetime.now().isoformat()
        }
        
        if metadata:
            payload.update(metadata)
        
        points.append(PointStruct(
            id=point_id,
            vector=embedding,
            payload=payload
        ))
    
    # Вставляем в Qdrant
    try:
        client.upsert(
            collection_name=COLLECTION_NAME,
            points=points
        )
        log.info(f"✅ Индексировано {len(points)} чанков из {file_path.name}")
        return True
    except Exception as e:
        log.error(f"❌ Ошибка индексации в Qdrant: {e}")
        import traceback
        log.error(f"❌ Traceback: {traceback.format_exc()}")
        return False

def index_directory(
    directory: Path,
    category: str = "другое",
    recursive: bool = True
) -> Dict:
    """
    Индексирование всех документов в директории
    
    Args:
        directory: Путь к директории
        category: Категория документов
        recursive: Рекурсивный поиск
    
    Returns:
        Статистика индексации
    """
    stats = {
        "total": 0,
        "success": 0,
        "failed": 0,
        "files": []
    }
    
    # Поддерживаемые форматы
    extensions = [".docx", ".xlsx", ".xls", ".pdf"]
    
    pattern = "**/*" if recursive else "*"
    for file_path in directory.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            stats["total"] += 1
            log.info(f"📄 Индексирование: {file_path}")
            
            if index_document(file_path, category=category):
                stats["success"] += 1
                stats["files"].append(str(file_path))
            else:
                stats["failed"] += 1
    
    return stats

# ===================== MAIN =====================

def main():
    """Главная функция скрипта"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Индексация базы знаний в Qdrant")
    parser.add_argument("path", type=str, help="Путь к файлу или директории")
    parser.add_argument("--category", type=str, default="другое", 
                       help="Категория документов (кейсы, методики, шаблоны, статьи)")
    parser.add_argument("--recursive", action="store_true", 
                       help="Рекурсивный поиск в директории")
    
    args = parser.parse_args()
    
    path = Path(args.path)
    if not path.exists():
        log.error(f"❌ Путь не существует: {path}")
        sys.exit(1)
    
    if path.is_file():
        log.info(f"📄 Индексация файла: {path}")
        success = index_document(path, category=args.category)
        sys.exit(0 if success else 1)
    elif path.is_dir():
        log.info(f"📁 Индексация директории: {path}")
        stats = index_directory(path, category=args.category, recursive=args.recursive)
        
        log.info("\n" + "="*50)
        log.info("📊 Статистика индексации:")
        log.info(f"  Всего файлов: {stats['total']}")
        log.info(f"  Успешно: {stats['success']}")
        log.info(f"  Ошибок: {stats['failed']}")
        log.info("="*50)
        
        sys.exit(0 if stats['failed'] == 0 else 1)
    else:
        log.error(f"❌ Некорректный путь: {path}")
        sys.exit(1)

if __name__ == "__main__":
    main()


